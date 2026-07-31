from __future__ import annotations

import re
import subprocess
import time
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
ARTICLE = ROOT / "content" / "posts" / "2026-06-11-llm-will-not-replace-jobs.md"
DESKTOP = Path.home() / "Desktop"
DOCX_OUT = DESKTOP / "LLM永远不会替代任何岗位，更不会替代人类.docx"
PDF_OUT = DESKTOP / "LLM永远不会替代任何岗位，更不会替代人类.pdf"
ARTICLE_URL = "http://127.0.0.1:1313/posts/2026-06-11-llm-will-not-replace-jobs/"


def split_front_matter(text: str) -> tuple[dict, str]:
    match = re.match(r"^---\s*\n(.*?)\n---\s*\n(.*)$", text, re.S)
    if not match:
        return {}, text
    return yaml.safe_load(match.group(1)) or {}, match.group(2).strip()


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_markdown_runs(paragraph, text: str, size: float = 11.5) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(64)
    section.bottom_margin = Pt(64)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_after = Pt(8)

    for name, size in [("Heading 1", 18), ("Heading 2", 15)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(18)
        style.paragraph_format.space_after = Pt(8)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "D9D9D9")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def build_docx() -> Path:
    meta, body = split_front_matter(ARTICLE.read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)

    title = str(meta.get("title") or "LLM 永远不会替代任何岗位，更不会替代人类")
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(10)
    title_run = title_para.add_run(title)
    set_run_font(title_run, size=22, bold=True)

    tags = " / ".join(str(tag) for tag in meta.get("tags", []))
    date = str(meta.get("date", ""))[:10]
    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(16)
    meta_run = meta_para.add_run(f"{date}  ·  {tags}")
    set_run_font(meta_run, size=9.5)
    add_bottom_border(meta_para)

    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            p = doc.add_heading(block[3:].strip(), level=1)
            set_paragraph_font(p, size=16)
            continue
        p = doc.add_paragraph()
        add_markdown_runs(p, block, size=11.5)

    footer = doc.sections[0].footer.paragraphs[0]
    footer_run = footer.add_run("Generated from local Hugo blog")
    set_run_font(footer_run, size=8.5)

    doc.save(DOCX_OUT)
    return DOCX_OUT


def ensure_hugo_server() -> None:
    import urllib.request

    try:
        urllib.request.urlopen(ARTICLE_URL, timeout=2).read(1)
        return
    except Exception:
        pass

    hugo_candidates = list(
        (Path.home() / "AppData" / "Local" / "Microsoft" / "WinGet" / "Packages").rglob("hugo.exe")
    )
    if not hugo_candidates:
        raise RuntimeError("hugo.exe not found")
    subprocess.Popen(
        [str(hugo_candidates[0]), "server", "-D", "--bind", "127.0.0.1", "--port", "1313"],
        cwd=str(ROOT),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        creationflags=subprocess.CREATE_NO_WINDOW,
    )
    for _ in range(20):
        try:
            urllib.request.urlopen(ARTICLE_URL, timeout=1).read(1)
            return
        except Exception:
            time.sleep(0.5)
    raise RuntimeError("Hugo server did not start")


def export_pdf() -> Path:
    ensure_hugo_server()
    chrome = Path("C:/Program Files/Google/Chrome/Application/chrome.exe")
    if not chrome.exists():
        raise RuntimeError("Chrome not found")
    subprocess.run(
        [
            str(chrome),
            "--headless",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--print-to-pdf={PDF_OUT}",
            ARTICLE_URL,
        ],
        check=True,
        cwd=str(ROOT),
    )
    return PDF_OUT


def main() -> None:
    DESKTOP.mkdir(parents=True, exist_ok=True)
    docx = build_docx()
    pdf = export_pdf()
    print(docx)
    print(pdf)


if __name__ == "__main__":
    main()
